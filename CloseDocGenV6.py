"""
Project Closure Document Generator (v5)
---------------------------------------
New in v5 (UI overhaul):
- Tabbed layout: ✏️ Fill In / 👀 Preview / 📤 Export — no more endless scrolling
- Live sidebar checklist with progress bar showing which sections are done
  and which required ones are missing
- Generate button is disabled until requirements are met (with the blocking
  items listed next to it) instead of erroring after the click
- Generation shows step-by-step status; dark-mode-safe styling (heading color
  override is now scoped to light theme only)

Carried over from v4:
- Save/Load drafts (JSON), Document Control with version history,
  Budget with auto variance, PDF export with Unicode font handling (₹),
  single source of truth for sections, auto numbering, Word TOC, page numbers.
"""

import json
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from xml.sax.saxutils import escape

import pandas as pd
import streamlit as st
from docx import Document
from ai_success import assess_project_success, DEFAULT_WEIGHTS, CRITERIA
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Flowable,
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table as RLTable,
    TableStyle,
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"
BRAND_COLOR = "#1f4e79"
DRAFT_SCHEMA_VERSION = 1


def _register_pdf_fonts():
    """reportlab's built-in Helvetica is Latin-1 only — the ₹ glyph renders as
    a blank box. Prefer DejaVu Sans (present on most Linux servers and
    Streamlit Cloud); fall back to Helvetica + currency transliteration."""
    candidates = [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
    ]
    for regular, bold in candidates:
        try:
            pdfmetrics.registerFont(TTFont("AppSans", regular))
            pdfmetrics.registerFont(TTFont("AppSans-Bold", bold))
            return "AppSans", "AppSans-Bold", True
        except Exception:  # noqa: BLE001 - try the next candidate
            continue
    return "Helvetica", "Helvetica-Bold", False


PDF_FONT, PDF_FONT_BOLD, PDF_UNICODE_OK = _register_pdf_fonts()


def pdf_safe(text):
    """Transliterate glyphs Helvetica can't draw when no Unicode font exists."""
    if PDF_UNICODE_OK:
        return text
    return text.replace("₹", "INR ")


# ---------------- PAGE CONFIG ---------------- #
st.set_page_config(page_title="Closure Document Generator", page_icon="📄", layout="wide")

st.markdown(
    """
    <style>
    .block-container {padding-top: 2rem;}
    /* Brand heading color only on light theme — #1f4e79 is unreadable on dark */
    @media (prefers-color-scheme: light) {
        h1, h2, h3 {color: #1f4e79;}
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------- SECTION CONFIG (single source of truth) ---------------- #
@dataclass
class Section:
    key: str            # session_state / data dict key
    title: str          # heading text (without number)
    required: bool = False


CORE_SECTIONS = [
    Section("overview", "Project Overview", required=True),
    Section("objectives", "Objectives", required=True),
    Section("timeline", "Timeline"),
    Section("budget", "Budget Summary"),
    Section("deliverables", "Deliverables"),
    Section("risks", "Risks & Issues"),
    Section("lessons", "Lessons Learned"),
    Section("stakeholders", "Stakeholders"),
]

PROJECT_STATUSES = ["Completed", "Completed with deviations", "Cancelled", "On hold"]
CURRENCIES = ["₹ (INR)", "$ (USD)", "€ (EUR)", "£ (GBP)", "kr (DKK)"]

# Plain-text fields included in drafts (widget key -> default)
DRAFT_TEXT_FIELDS = {
    "project_name": "",
    "prepared_by": "",
    "overview": "",
    "objectives": "",
    "budget_notes": "",
    "deliverables": "",
    "risks": "",
    "lessons": "",
    "doc_version": "1.0",
}

EMPTY_STAKEHOLDERS = pd.DataFrame([{"Name": "", "Role": ""}])
EMPTY_HISTORY = pd.DataFrame(
    [{"Version": "1.0", "Date": date.today().strftime("%d %b %Y"),
      "Author": "", "Change Description": "Initial version"}]
)


# ---------------- SESSION STATE DEFAULTS ---------------- #
def init_state():
    defaults = {
        "custom_sections": [],
        "custom_section_seq": 0,
        "stakeholders_df": EMPTY_STAKEHOLDERS.copy(),
        "history_df": EMPTY_HISTORY.copy(),
        "editor_nonce": 0,   # bumped on draft load to force data_editor refresh
        "draft_loaded_token": None,
        "success_result": None,                  # cached AI assessment (dict)
        "success_weights": dict(DEFAULT_WEIGHTS), # user-adjustable criteria weights
        "success_model": "gpt-4o-mini",
        "include_score_in_export": True,
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


init_state()


# ---------------- DRAFT SAVE / LOAD ---------------- #
def df_to_rows(df, columns):
    rows = []
    for _, row in df.iterrows():
        values = [str(row.get(col, "") or "").strip() for col in columns]
        if any(values):
            rows.append(values)
    return rows


def build_draft(values, stakeholder_rows, history_rows):
    return {
        "schema": DRAFT_SCHEMA_VERSION,
        "app": "closure-doc-generator",
        "saved_at": datetime.now().isoformat(timespec="seconds"),
        "fields": {k: values[k] for k in DRAFT_TEXT_FIELDS},
        "status": values["status"],
        "currency": values["currency"],
        "planned_budget": values["planned_budget"],
        "actual_spend": values["actual_spend"],
        "start_date": values["start_date"].isoformat() if values["start_date"] else None,
        "end_date": values["end_date"].isoformat() if values["end_date"] else None,
        "stakeholders": stakeholder_rows,
        "version_history": history_rows,
        "custom_sections": [
            {"title": s["title"], "content": s["content"]}
            for s in st.session_state.custom_sections
        ],
        "success_result": st.session_state.get("success_result"),
        "success_weights": st.session_state.get("success_weights"),
    }


def apply_draft(draft):
    """Write draft values into session_state. MUST run before any of the
    target widgets are instantiated in this script run."""
    if draft.get("app") != "closure-doc-generator":
        raise ValueError("This JSON file was not saved by this app.")

    fields = draft.get("fields", {})
    for key, default in DRAFT_TEXT_FIELDS.items():
        st.session_state[key] = str(fields.get(key, default) or default)

    if draft.get("status") in PROJECT_STATUSES:
        st.session_state["status"] = draft["status"]
    if draft.get("currency") in CURRENCIES:
        st.session_state["currency"] = draft["currency"]

    st.session_state["planned_budget"] = float(draft.get("planned_budget") or 0.0)
    st.session_state["actual_spend"] = float(draft.get("actual_spend") or 0.0)

    for key in ("start_date", "end_date"):
        raw = draft.get(key)
        st.session_state[key] = date.fromisoformat(raw) if raw else None

    rows = draft.get("stakeholders") or []
    st.session_state.stakeholders_df = (
        pd.DataFrame(rows, columns=["Name", "Role"]) if rows else EMPTY_STAKEHOLDERS.copy()
    )
    rows = draft.get("version_history") or []
    st.session_state.history_df = (
        pd.DataFrame(rows, columns=["Version", "Date", "Author", "Change Description"])
        if rows else EMPTY_HISTORY.copy()
    )

    st.session_state.custom_sections = []
    st.session_state.custom_section_seq = 0
    for sec in draft.get("custom_sections", []):
        st.session_state.custom_section_seq += 1
        sid = st.session_state.custom_section_seq
        st.session_state.custom_sections.append(
            {"id": sid, "title": sec.get("title", ""), "content": sec.get("content", "")}
        )
        st.session_state[f"cs_title_{sid}"] = sec.get("title", "")
        st.session_state[f"cs_content_{sid}"] = sec.get("content", "")

    st.session_state.editor_nonce += 1  # force data_editors to pick up new frames

    # Restore a previously computed success assessment, if present
    if draft.get("success_result"):
        st.session_state.success_result = draft["success_result"]
    if draft.get("success_weights"):
        st.session_state.success_weights = draft["success_weights"]
        for c in CRITERIA:
            st.session_state.pop(f"weight_{c.key}", None)


# ---------------- BUDGET HELPERS ---------------- #
def money(amount, symbol):
    return f"{symbol}{amount:,.2f}"


def budget_rows(budget):
    """Shared Planned/Actual/Variance rows for DOCX, PDF, and preview."""
    symbol = budget["currency"].split(" ")[0]
    planned, actual = budget["planned"], budget["actual"]
    variance = planned - actual
    pct = (variance / planned * 100) if planned else 0.0
    verdict = "Under budget" if variance >= 0 else "Over budget"
    return [
        ["Planned Budget", money(planned, symbol)],
        ["Actual Spend", money(actual, symbol)],
        ["Variance", f"{money(variance, symbol)} ({pct:+.1f}%) — {verdict}"],
    ]


# ---------------- SUCCESS SCORE RENDER HELPERS ---------------- #
def success_summary_rows(result, include_emoji=True):
    """Top summary rows shared by preview/DOCX/PDF.
    PDF/DOCX fonts lack colour-emoji glyphs, so callers pass include_emoji=False."""
    band = f"{result['band_emoji']} {result['classification']}" if include_emoji \
        else result["classification"]
    return [
        ["Final Success Score", f"{result['final_score']} / 100"],
        ["Classification", band],
        ["Scoring Engine", result["model"]],
    ]


def success_breakdown_rows(result):
    """Per-criterion table: Criterion | Weight | Score | Weighted | Justification."""
    rows = [["Criterion", "Weight", "Score", "Weighted", "Justification"]]
    for c in result["criteria"]:
        weighted = round(c["score"] * c["weight"], 1)
        rows.append([
            c["label"],
            f"{c['weight'] * 100:.0f}%",
            f"{c['score']:.0f}",
            f"{weighted:.1f}",
            c["justification"],
        ])
    return rows


def docx_render_success(doc, result):
    add_table(doc, [["Field", "Value"]] + success_summary_rows(result, include_emoji=False))
    doc.add_paragraph()
    add_table(doc, success_breakdown_rows(result))
    if result.get("engine") == "fallback":
        doc.add_paragraph()
        add_paragraph(doc, f"Note: {result.get('notes', '')}")


# ---------------- DOCX HELPERS ---------------- #
def add_paragraph(doc, text):
    para = doc.add_paragraph(text.strip() if text and text.strip() else "N/A")
    para.paragraph_format.space_after = Pt(10)


def add_table(doc, data, style="Light Grid Accent 1"):
    """data: list of rows; first row is the header."""
    table = doc.add_table(rows=1, cols=len(data[0]))
    try:
        table.style = style
    except KeyError:
        table.style = "Table Grid"  # fallback if the style isn't in the template
    for i, heading in enumerate(data[0]):
        cell = table.rows[0].cells[i]
        cell.text = str(heading)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in data[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def add_table_of_contents(doc):
    """Insert a real Word TOC field. Word shows placeholder text until the
    user updates fields (Ctrl+A, F9) — the 'dirty' flag makes Word offer this."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")  # ask Word to refresh on open

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to generate the Table of Contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_separate, placeholder, fld_end):
        run._r.append(el)


def add_page_numbers(doc):
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_end):
        run._r.append(el)


def sanitize_filename(name):
    cleaned = "".join(c for c in name if c.isalnum() or c in (" ", "_", "-")).strip()
    return cleaned or "Project"


# ---------------- DOCX SECTION RENDERERS ---------------- #
def docx_render_budget(doc, data):
    budget = data["budget"]
    if not budget["planned"] and not budget["actual"]:
        add_paragraph(doc, "N/A")
        return
    add_table(doc, [["Item", "Amount"]] + budget_rows(budget))
    if budget["notes"].strip():
        doc.add_paragraph()
        add_paragraph(doc, budget["notes"])


def docx_render_deliverables(doc, data):
    add_paragraph(doc, data["deliverables"])
    files = data["files"]
    if not files:
        return

    doc.add_paragraph("Attached Deliverables:")
    table_data = [["File Name", "Type"]]
    for file in files:
        try:
            if file.type and file.type.startswith("image"):
                doc.add_paragraph(f"Image: {file.name}")
                file.seek(0)
                doc.add_picture(BytesIO(file.read()), width=Inches(5))
            else:
                table_data.append([file.name, file.type or "Unknown"])
        except Exception as exc:  # noqa: BLE001 - keep generating, note the failure
            table_data.append([file.name, f"Could not embed ({exc.__class__.__name__})"])
    if len(table_data) > 1:
        add_table(doc, table_data)


def docx_render_stakeholders(doc, data):
    rows = data["stakeholders"] or [["N/A", "N/A"]]
    add_table(doc, [["Name", "Role"]] + rows)


def generate_docx(data):
    doc = Document()

    # Metadata (visible in File > Info, useful for SharePoint/search)
    props = doc.core_properties
    props.title = f"{data['project_name']} – Project Closure Report"
    props.author = data["prepared_by"] or "Project Manager"
    props.subject = "Project Closure Report"
    props.version = data["doc_version"]
    props.created = datetime.now()

    # Cover page
    doc.add_heading(data["project_name"], 0)
    add_paragraph(doc, "Project Closure Report")
    cover = [
        ["Document Version", data["doc_version"]],
        ["Project Status", data["status"]],
        ["Timeline", data["timeline"] or "N/A"],
        ["Prepared By", data["prepared_by"] or "N/A"],
        ["Generated On", datetime.now().strftime("%d %b %Y")],
    ]
    add_table(doc, [["Field", "Detail"]] + cover)

    # Document control (standard opener for formal closure documents)
    doc.add_paragraph()
    doc.add_heading("Document Control", 1)
    history = data["version_history"] or [["N/A", "N/A", "N/A", "N/A"]]
    add_table(doc, [["Version", "Date", "Author", "Change Description"]] + history)
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", 1)
    add_table_of_contents(doc)
    doc.add_page_break()

    # Numbered sections — numbering derived from position, never hardcoded
    renderers = {
        "budget": docx_render_budget,
        "deliverables": docx_render_deliverables,
        "stakeholders": docx_render_stakeholders,
    }
    number = 0
    for section in CORE_SECTIONS:
        number += 1
        doc.add_heading(f"{number}. {section.title}", 1)
        renderer = renderers.get(section.key)
        if renderer:
            renderer(doc, data)
        else:
            add_paragraph(doc, data[section.key])

    # Success assessment (optional, inserted after core sections)
    success = data.get("success_result")
    if success and data.get("include_score"):
        number += 1
        doc.add_heading(f"{number}. Project Success Assessment", 1)
        docx_render_success(doc, success)

    # Custom sections continue the same numbering
    for sec in data["custom_sections"]:
        if sec["title"].strip():
            number += 1
            doc.add_heading(f"{number}. {sec['title'].strip()}", 1)
            add_paragraph(doc, sec["content"])

    # Sign-off
    number += 1
    doc.add_heading(f"{number}. Approval & Sign-off", 1)
    add_table(doc, [["Name", "Role", "Signature", "Date"], ["", "", "", ""], ["", "", "", ""]])

    add_page_numbers(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------- PDF GENERATION (reportlab) ---------------- #
def _pdf_styles():
    base = getSampleStyleSheet()
    brand = colors.HexColor(BRAND_COLOR)
    return {
        "title": ParagraphStyle("DocTitle", parent=base["Title"], textColor=brand,
                                fontSize=24, fontName=PDF_FONT_BOLD),
        "subtitle": ParagraphStyle("Subtitle", parent=base["Normal"], fontSize=12,
                                   textColor=colors.grey, spaceAfter=18, fontName=PDF_FONT),
        "h1": ParagraphStyle("H1", parent=base["Heading1"], textColor=brand,
                             spaceBefore=16, spaceAfter=8, fontName=PDF_FONT_BOLD),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontSize=10.5,
                               leading=15, spaceAfter=8, fontName=PDF_FONT),
    }


def _pdf_table(rows, col_widths=None):
    brand = colors.HexColor(BRAND_COLOR)
    rows = [[cell if isinstance(cell, Flowable) else pdf_safe(str(cell))
             for cell in row] for row in rows]
    table = RLTable(rows, colWidths=col_widths, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), brand),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), PDF_FONT),
        ("FONTNAME", (0, 0), (-1, 0), PDF_FONT_BOLD),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B7C6D9")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF3F8")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return table


def _pdf_text(text, style):
    """Multiline plain text -> Paragraph (escaped, newlines preserved)."""
    content = text.strip() if text and text.strip() else "N/A"
    return Paragraph(escape(pdf_safe(content)).replace("\n", "<br/>"), style)


def _pdf_image(file, max_width):
    file.seek(0)
    raw = file.read()
    img_w, img_h = ImageReader(BytesIO(raw)).getSize()
    width = min(max_width, img_w)
    return RLImage(BytesIO(raw), width=width, height=img_h * width / img_w)


def _pdf_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont(PDF_FONT, 9)
    canvas.setFillColor(colors.grey)
    canvas.drawCentredString(A4[0] / 2, 1 * cm, str(canvas.getPageNumber()))
    canvas.restoreState()


def generate_pdf(data):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
        title=f"{data['project_name']} – Project Closure Report",
        author=data["prepared_by"] or "Project Manager",
    )
    s = _pdf_styles()
    content_width = A4[0] - 4 * cm
    story = []

    # Cover
    story.append(Paragraph(escape(data["project_name"]), s["title"]))
    story.append(Paragraph("Project Closure Report", s["subtitle"]))
    cover = [
        ["Field", "Detail"],
        ["Document Version", data["doc_version"]],
        ["Project Status", data["status"]],
        ["Timeline", data["timeline"] or "N/A"],
        ["Prepared By", data["prepared_by"] or "N/A"],
        ["Generated On", datetime.now().strftime("%d %b %Y")],
    ]
    story.append(_pdf_table(cover, [content_width * 0.35, content_width * 0.65]))

    # Document control
    story.append(Paragraph("Document Control", s["h1"]))
    history = data["version_history"] or [["N/A", "N/A", "N/A", "N/A"]]
    story.append(_pdf_table(
        [["Version", "Date", "Author", "Change Description"]] + history,
        [content_width * 0.12, content_width * 0.18, content_width * 0.25, content_width * 0.45],
    ))
    story.append(PageBreak())

    def add_section(title, flowables):
        story.append(Paragraph(escape(title), s["h1"]))
        story.extend(flowables)
        story.append(Spacer(1, 6))

    number = 0
    for section in CORE_SECTIONS:
        number += 1
        title = f"{number}. {section.title}"

        if section.key == "budget":
            budget = data["budget"]
            if not budget["planned"] and not budget["actual"]:
                add_section(title, [_pdf_text("", s["body"])])
            else:
                flowables = [_pdf_table([["Item", "Amount"]] + budget_rows(budget),
                                        [content_width * 0.35, content_width * 0.65])]
                if budget["notes"].strip():
                    flowables += [Spacer(1, 8), _pdf_text(budget["notes"], s["body"])]
                add_section(title, flowables)

        elif section.key == "deliverables":
            flowables = [_pdf_text(data["deliverables"], s["body"])]
            files = data["files"] or []
            non_images = [["File Name", "Type"]]
            for f in files:
                try:
                    if f.type and f.type.startswith("image"):
                        flowables += [Spacer(1, 6), _pdf_text(f"Image: {f.name}", s["body"]),
                                      _pdf_image(f, content_width)]
                    else:
                        non_images.append([f.name, f.type or "Unknown"])
                except Exception as exc:  # noqa: BLE001
                    non_images.append([f.name, f"Could not embed ({exc.__class__.__name__})"])
            if len(non_images) > 1:
                flowables += [Spacer(1, 8), _pdf_table(non_images)]
            add_section(title, flowables)

        elif section.key == "stakeholders":
            rows = data["stakeholders"] or [["N/A", "N/A"]]
            add_section(title, [_pdf_table([["Name", "Role"]] + rows,
                                           [content_width * 0.5, content_width * 0.5])])
        else:
            add_section(title, [_pdf_text(data[section.key], s["body"])])

    success = data.get("success_result")
    if success and data.get("include_score"):
        number += 1
        summary_tbl = _pdf_table(
            [["Field", "Value"]] + success_summary_rows(success, include_emoji=False),
            [content_width * 0.35, content_width * 0.65],
        )
        breakdown = success_breakdown_rows(success)
        # Wrap the criterion label and justification columns so text doesn't overlap.
        wrap_style = ParagraphStyle("CellWrap", parent=s["body"], fontSize=8.5, leading=11)
        for r in range(1, len(breakdown)):
            breakdown[r][0] = Paragraph(escape(pdf_safe(str(breakdown[r][0]))), wrap_style)
            breakdown[r][-1] = Paragraph(escape(pdf_safe(str(breakdown[r][-1]))), wrap_style)
        breakdown_tbl = _pdf_table(
            breakdown,
            [content_width * 0.24, content_width * 0.10, content_width * 0.09,
             content_width * 0.12, content_width * 0.45],
        )
        flowables = [summary_tbl, Spacer(1, 10), breakdown_tbl]
        if success.get("engine") == "fallback":
            flowables += [Spacer(1, 6), _pdf_text(f"Note: {success.get('notes','')}", s["body"])]
        add_section(f"{number}. Project Success Assessment", flowables)

    for sec in data["custom_sections"]:
        if sec["title"].strip():
            number += 1
            add_section(f"{number}. {sec['title'].strip()}", [_pdf_text(sec["content"], s["body"])])

    number += 1
    signoff = [["Name", "Role", "Signature", "Date"], ["", "", "", ""], ["", "", "", ""]]
    add_section(f"{number}. Approval & Sign-off", [_pdf_table(signoff, [content_width / 4] * 4)])

    pdf.build(story, onFirstPage=_pdf_page_number, onLaterPages=_pdf_page_number)
    buffer.seek(0)
    return buffer


# ---------------- UI ---------------- #
st.title("📄 Project Closure Generator")

tab_fill, tab_score, tab_preview, tab_export = st.tabs(
    ["✏️ Fill In", "🎯 Success Score", "👀 Preview", "📤 Export"]
)

# ============================================================== #
# TAB 1 — FILL IN
# (Draft loader MUST be the first block: Streamlit forbids writing
#  to a widget's session_state key after that widget has rendered.)
# ============================================================== #
with tab_fill:
    with st.expander("💾 Load a Draft"):
        st.caption(
            "👉 Restore a previously saved draft (JSON). "
            "Uploaded deliverable files aren't included in drafts — re-attach them after loading."
        )
        draft_file = st.file_uploader("Load a draft", type=["json"], key="draft_uploader",
                                      label_visibility="collapsed")
        if draft_file is not None:
            token = (draft_file.name, draft_file.size)
            already_loaded = st.session_state.draft_loaded_token == token
            if st.button("📂 Load this draft", disabled=already_loaded,
                         help="Already loaded" if already_loaded else None):
                try:
                    draft = json.loads(draft_file.getvalue().decode("utf-8"))
                    apply_draft(draft)
                    st.session_state.draft_loaded_token = token
                    st.rerun()
                except (ValueError, KeyError, json.JSONDecodeError) as exc:
                    st.error(f"Could not load draft: {exc}")
            if already_loaded:
                st.success("Draft loaded. Remove the file above to load a different one.")

    with st.expander("🗂️ Document Control"):
        st.caption("👉 Version history — required in most formal closure documents")
        doc_version = st.text_input("Document Version", key="doc_version", placeholder="1.0")
        st.write("**Version History** (add rows as the document evolves)")
        history_df = st.data_editor(
            st.session_state.history_df,
            num_rows="dynamic",
            width="stretch",
            key=f"history_editor_{st.session_state.editor_nonce}",
        )

    with st.expander("📌 Project Information", expanded=True):
        st.caption("👉 Basic identification details")
        col1, col2 = st.columns(2)
        with col1:
            project_name = st.text_input(
                "Project Name *",
                key="project_name",
                help="Use official project name",
                placeholder="Migration of On-Prem Email to Cloud",
            )
            prepared_by = st.text_input("Prepared By", key="prepared_by", placeholder="Your name")
        with col2:
            status = st.selectbox("Project Status", PROJECT_STATUSES, key="status")
            date_col1, date_col2 = st.columns(2)
            start_date = date_col1.date_input("Start Date", value=None, key="start_date",
                                              max_value=date.today())
            end_date = date_col2.date_input("End Date", value=None, key="end_date")

        if start_date and end_date and end_date < start_date:
            st.error("End date cannot be before start date.")

        timeline = ""
        if start_date and end_date:
            timeline = f"{start_date.strftime('%d %b %Y')} – {end_date.strftime('%d %b %Y')}"

    with st.expander("📖 Project Overview", expanded=True):
        st.caption("👉 What was done, why, and outcome — cover scope + purpose + final result")
        overview = st.text_area(
            "Overview *",
            key="overview",
            help="Write a concise 3–5 line summary",
            placeholder="Migrated 2500+ users to cloud, improving availability and reducing infra dependency.",
        )

    with st.expander("🎯 Objectives"):
        st.caption("👉 What success looked like — use measurable goals")
        objectives = st.text_area(
            "Objectives *",
            key="objectives",
            placeholder="- Reduce downtime < 1 hour\n- Improve remote access\n- Decommission legacy infra",
        )

    with st.expander("💰 Budget"):
        st.caption("👉 Planned vs actual — variance is calculated for you; "
                   "explain anything over ±10% in the notes")
        bcol1, bcol2, bcol3 = st.columns([1, 1.5, 1.5])
        currency = bcol1.selectbox("Currency", CURRENCIES, key="currency")
        planned_budget = bcol2.number_input("Planned Budget", min_value=0.0, step=1000.0,
                                            format="%.2f", key="planned_budget")
        actual_spend = bcol3.number_input("Actual Spend", min_value=0.0, step=1000.0,
                                          format="%.2f", key="actual_spend")

        if planned_budget or actual_spend:
            symbol = currency.split(" ")[0]
            variance = planned_budget - actual_spend
            pct = (variance / planned_budget * 100) if planned_budget else 0.0
            st.metric(
                "Variance",
                money(variance, symbol),
                delta=f"{pct:+.1f}% ({'under' if variance >= 0 else 'over'} budget)",
            )

        budget_notes = st.text_area(
            "Budget Notes",
            key="budget_notes",
            placeholder="E.g., Overspend driven by extended UAT phase; offset by license savings.",
        )

    with st.expander("📦 Deliverables"):
        st.caption("👉 What was delivered — think outputs, not activities")
        deliverables = st.text_area(
            "Deliverables",
            key="deliverables",
            placeholder="- Mailboxes migrated\n- Documentation created\n- Dashboard deployed",
        )
        uploaded_files = st.file_uploader(
            "Upload Supporting Files",
            accept_multiple_files=True,
            type=["png", "jpg", "jpeg", "pdf", "docx", "xlsx", "csv", "txt"],
            help="Upload screenshots, reports, dashboards (not saved in drafts)",
        )

    with st.expander("⚠️ Risks & Issues"):
        st.caption("👉 What went wrong or could have — mention impact + mitigation")
        risks = st.text_area("Risks", key="risks",
                             placeholder="Data corruption risk, delays, etc.")

    with st.expander("📘 Lessons Learned"):
        st.caption("👉 What should improve next time — focus on actionable improvements")
        lessons = st.text_area("Lessons", key="lessons",
                               placeholder="Validate data earlier, improve planning buffer")

    with st.expander("👥 Stakeholders"):
        st.caption("👉 Who was involved — add rows as needed")
        stakeholder_df = st.data_editor(
            st.session_state.stakeholders_df,
            num_rows="dynamic",
            width="stretch",
            key=f"stakeholder_editor_{st.session_state.editor_nonce}",
        )

    # ---------------- CUSTOM SECTIONS ---------------- #
    st.subheader("➕ Custom Sections")

    if st.button("Add Section"):
        st.session_state.custom_section_seq += 1
        st.session_state.custom_sections.append(
            {"id": st.session_state.custom_section_seq, "title": "", "content": ""}
        )

    for sec in st.session_state.custom_sections:
        sid = sec["id"]
        title_col, remove_col = st.columns([6, 1])
        sec["title"] = title_col.text_input(
            "Section Title", key=f"cs_title_{sid}", placeholder="E.g., Migration Strategy"
        )
        if remove_col.button("🗑️", key=f"cs_remove_{sid}", help="Remove this section"):
            st.session_state.custom_sections = [
                s for s in st.session_state.custom_sections if s["id"] != sid
            ]
            st.rerun()
        sec["content"] = st.text_area(
            "Section Content", key=f"cs_content_{sid}", placeholder="Describe this section"
        )

# ---------------- DERIVED VALUES (used by preview, export, sidebar) ----------- #
stakeholders = df_to_rows(stakeholder_df, ["Name", "Role"])
version_history = df_to_rows(history_df, ["Version", "Date", "Author", "Change Description"])

budget_data = {
    "currency": currency,
    "planned": planned_budget,
    "actual": actual_spend,
    "notes": budget_notes,
}


# ============================================================== #
# TAB 2 — SUCCESS SCORE (AI)
# ============================================================== #
with tab_score:
    st.subheader("🎯 AI Project Success Assessment")
    st.caption(
        "Scores the project against weighted criteria. Budget and schedule are "
        "computed deterministically; objectives, deliverables, risk and quality "
        "are evaluated by AI. Needs an OpenAI key — set OPENAI_API_KEY in your "
        "local .env or in Streamlit Cloud → Settings → Secrets. Without a key it "
        "falls back to an offline heuristic."
    )

    cfg_col1, cfg_col2 = st.columns([1, 1])
    st.session_state.success_model = cfg_col1.selectbox(
        "AI Model",
        ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo"],
        index=["gpt-4o-mini", "gpt-4o", "gpt-4-turbo"].index(st.session_state.success_model),
        help="gpt-4o-mini is cheapest; gpt-4o is the most capable.",
    )
    st.session_state.include_score_in_export = cfg_col2.checkbox(
        "Include score section in exported Word/PDF",
        value=st.session_state.include_score_in_export,
    )

    with st.expander("⚖️ Adjust criteria weights"):
        st.caption("Weights are normalised automatically, so they don't need to sum to 100%.")
        wcols = st.columns(3)
        new_weights = {}
        for i, c in enumerate(CRITERIA):
            with wcols[i % 3]:
                new_weights[c.key] = st.slider(
                    c.label,
                    min_value=0, max_value=50,
                    value=int(round(st.session_state.success_weights.get(c.key, c.default_weight) * 100)),
                    step=5, key=f"weight_{c.key}",
                ) / 100.0
        st.session_state.success_weights = new_weights
        total = sum(new_weights.values())
        st.caption(f"Raw total: {total * 100:.0f}% (will be normalised to 100%).")
        if st.button("↩️ Reset to default weights"):
            st.session_state.success_weights = dict(DEFAULT_WEIGHTS)
            for c in CRITERIA:
                st.session_state.pop(f"weight_{c.key}", None)
            st.rerun()

    ready = bool(objectives.strip() and deliverables.strip())
    if not ready:
        st.info("Fill in **Objectives** and **Deliverables** (in the ✏️ Fill In tab) "
                "to get a meaningful score.")

    if st.button("🤖 Assess Project Success", type="primary",
                 disabled=not (objectives.strip() or deliverables.strip())):
        payload = {
            "project_name": project_name,
            "status": status,
            "overview": overview,
            "objectives": objectives,
            "deliverables": deliverables,
            "risks": risks,
            "lessons": lessons,
            "budget": {"planned": planned_budget, "actual": actual_spend, "currency": currency},
            "start_date": start_date,
            "end_date": end_date,
        }
        with st.spinner("Assessing project against success criteria..."):
            try:
                st.session_state.success_result = assess_project_success(
                    payload,
                    weights=st.session_state.success_weights,
                    model=st.session_state.success_model,
                )
            except Exception as exc:  # noqa: BLE001
                st.error(f"Assessment failed: {exc}")

    result = st.session_state.success_result
    if result:
        if result.get("engine") == "fallback":
            st.warning(result.get("notes", "Used offline heuristic scoring."))

        m1, m2, m3 = st.columns(3)
        m1.metric("Final Success Score", f"{result['final_score']} / 100")
        m2.metric("Classification", f"{result['band_emoji']} {result['classification']}")
        m3.metric("Engine", result["model"])

        st.progress(min(result["final_score"] / 100, 1.0))

        st.markdown("#### Criteria breakdown")
        breakdown_df = pd.DataFrame([
            {
                "Criterion": c["label"],
                "Weight": f"{c['weight'] * 100:.0f}%",
                "Score": c["score"],
                "Weighted": round(c["score"] * c["weight"], 1),
                "Justification": c["justification"],
            }
            for c in result["criteria"]
        ])
        st.dataframe(breakdown_df, width="stretch", hide_index=True)

        chart_df = pd.DataFrame(
            {"Score": [c["score"] for c in result["criteria"]]},
            index=[c["label"] for c in result["criteria"]],
        )
        st.bar_chart(chart_df, horizontal=True)

        st.caption(f"Assessed at {result.get('assessed_at', 'N/A')}. "
                   "Re-run after editing fields to refresh.")


def validate():
    """Collect ALL problems at once. Used to disable Generate, not scold after."""
    errors = []
    if not project_name.strip():
        errors.append("Project Name is required.")
    if not overview.strip():
        errors.append("Overview is required.")
    if not objectives.strip():
        errors.append("Objectives are required.")
    if start_date and end_date and end_date < start_date:
        errors.append("End date cannot be before start date.")
    if actual_spend and not planned_budget:
        errors.append("Actual Spend entered without a Planned Budget — variance can't be computed.")
    for sec in st.session_state.custom_sections:
        if sec["content"].strip() and not sec["title"].strip():
            errors.append("A custom section has content but no title.")
    return errors


validation_errors = validate()

# Sidebar checklist items: (label, complete?, required?)
CHECKLIST = [
    ("Project Name", bool(project_name.strip()), True),
    ("Overview", bool(overview.strip()), True),
    ("Objectives", bool(objectives.strip()), True),
    ("Timeline", bool(timeline), False),
    ("Budget", bool(planned_budget or actual_spend), False),
    ("Deliverables", bool(deliverables.strip() or uploaded_files), False),
    ("Risks & Issues", bool(risks.strip()), False),
    ("Lessons Learned", bool(lessons.strip()), False),
    ("Stakeholders", bool(stakeholders), False),
]

# ============================================================== #
# TAB 2 — PREVIEW
# ============================================================== #
with tab_preview:
    if project_name:
        st.markdown(f"# {project_name}")
    else:
        st.info("Start filling in the ✏️ Fill In tab — the preview updates live.")

    st.markdown(
        f"**Version:** {doc_version or 'N/A'} &nbsp;|&nbsp; "
        f"**Status:** {status} &nbsp;|&nbsp; **Timeline:** {timeline or 'N/A'}"
    )

    with st.expander("Document Control"):
        if version_history:
            st.table(pd.DataFrame(version_history,
                                  columns=["Version", "Date", "Author", "Change Description"]))
        else:
            st.write("N/A")

    preview_values = {
        "overview": overview,
        "objectives": objectives,
        "timeline": timeline,
        "deliverables": deliverables,
        "risks": risks,
        "lessons": lessons,
    }

    number = 0
    for section in CORE_SECTIONS:
        number += 1
        with st.expander(f"{number}. {section.title}", expanded=(section.key == "overview")):
            if section.key == "budget":
                if planned_budget or actual_spend:
                    for label, value in budget_rows(budget_data):
                        st.write(f"**{label}:** {value}")
                    if budget_notes.strip():
                        st.write(budget_notes)
                else:
                    st.write("N/A")
            elif section.key == "stakeholders":
                if stakeholders:
                    for name, role in stakeholders:
                        st.write(f"- {name} ({role})")
                else:
                    st.write("N/A")
            elif section.key == "deliverables":
                st.write(deliverables or "N/A")
                if uploaded_files:
                    st.markdown("**Attached Files:**")
                    for f in uploaded_files:
                        if f.type and f.type.startswith("image"):
                            st.image(f, caption=f.name)
                        else:
                            st.write(f"📄 {f.name}")
            else:
                st.write(preview_values[section.key] or "N/A")

    for sec in st.session_state.custom_sections:
        if sec["title"].strip():
            number += 1
            with st.expander(f"{number}. {sec['title'].strip()}"):
                st.write(sec["content"] or "N/A")

    _score = st.session_state.success_result
    if _score and st.session_state.include_score_in_export:
        number += 1
        with st.expander(f"{number}. Project Success Assessment", expanded=True):
            st.write(f"**Final Score:** {_score['final_score']} / 100 — "
                     f"{_score['band_emoji']} {_score['classification']}")
            for c in _score["criteria"]:
                st.write(f"- **{c['label']}** ({c['weight'] * 100:.0f}%): "
                         f"{c['score']:.0f} — {c['justification']}")

# ============================================================== #
# TAB 3 — EXPORT
# ============================================================== #
with tab_export:
    st.subheader("💾 Save Draft")
    st.caption("Saves all text fields, dates, budget, tables, and custom sections as JSON. "
               "Uploaded files are not included.")
    draft_values = {
        **{k: st.session_state.get(k, v) for k, v in DRAFT_TEXT_FIELDS.items()},
        "status": status,
        "currency": currency,
        "planned_budget": planned_budget,
        "actual_spend": actual_spend,
        "start_date": start_date,
        "end_date": end_date,
    }
    st.download_button(
        "💾 Save Draft (JSON)",
        json.dumps(build_draft(draft_values, stakeholders, version_history),
                   indent=2, ensure_ascii=False),
        file_name=f"{sanitize_filename(project_name)}_closure_draft.json",
        mime="application/json",
    )

    st.divider()
    st.subheader("🚀 Generate Documents")

    can_generate = not validation_errors
    if not can_generate:
        st.caption("Complete the following to enable generation:")
        for err in validation_errors:
            st.caption(f"• {err}")

    if st.button("🚀 Generate Documents", type="primary", disabled=not can_generate,
                 help=None if can_generate else "Fix the items listed above first"):
        data = {
            "project_name": project_name.strip(),
            "prepared_by": prepared_by.strip(),
            "status": status,
            "doc_version": doc_version.strip() or "1.0",
            "version_history": version_history,
            "overview": overview,
            "objectives": objectives,
            "timeline": timeline,
            "budget": budget_data,
            "deliverables": deliverables,
            "files": uploaded_files,
            "risks": risks,
            "lessons": lessons,
            "stakeholders": stakeholders,
            "custom_sections": st.session_state.custom_sections,
            "success_result": st.session_state.success_result,
            "include_score": st.session_state.include_score_in_export,
        }

        try:
            with st.status("Generating documents...", expanded=True) as gen_status:
                st.write("Building Word document...")
                docx_buffer = generate_docx(data)
                st.write("Building PDF...")
                pdf_buffer = generate_pdf(data)
                gen_status.update(label="Documents ready ✅", state="complete", expanded=False)
        except Exception as exc:  # noqa: BLE001
            st.error(f"Document generation failed: {exc}")
            st.stop()

        # Persist across reruns: clicking one download button reruns the script,
        # which would otherwise discard the other file.
        st.session_state.generated = {
            "docx": docx_buffer.getvalue(),
            "pdf": pdf_buffer.getvalue(),
            "base_name": sanitize_filename(project_name),
        }

    if "generated" in st.session_state:
        gen = st.session_state.generated
        dl_col1, dl_col2 = st.columns(2)
        dl_col1.download_button(
            "📥 Download Word (.docx)",
            gen["docx"],
            file_name=f"{gen['base_name']}_Closure_Report.docx",
            mime=DOCX_MIME,
            width="stretch",
        )
        dl_col2.download_button(
            "📥 Download PDF (.pdf)",
            gen["pdf"],
            file_name=f"{gen['base_name']}_Closure_Report.pdf",
            mime=PDF_MIME,
            width="stretch",
        )

# ============================================================== #
# SIDEBAR — LIVE COMPLETION CHECKLIST
# (Rendered last in code so it reflects this run's values, but
#  Streamlit displays it in the sidebar regardless of code order.)
# ============================================================== #
with st.sidebar:
    st.header("📋 Progress")
    done = sum(1 for _, complete, _ in CHECKLIST if complete)
    st.progress(done / len(CHECKLIST), text=f"{done} of {len(CHECKLIST)} sections complete")

    for label, complete, required in CHECKLIST:
        icon = "✅" if complete else "⬜"
        suffix = "" if complete or not required else " *(required)*"
        st.markdown(f"{icon} {label}{suffix}")

    st.divider()
    if validation_errors:
        st.warning(f"{len(validation_errors)} item(s) blocking generation — see 📤 Export tab.")
    else:
        st.success("Ready to generate! Head to the 📤 Export tab.")

    with st.expander("📘 Writing Tips"):
        st.markdown(
            "- Write outcomes, not activities\n"
            "- Use numbers (%, count, time saved)\n"
            "- Keep it concise — decision makers skim\n"
            "- Add proof via files/screenshots\n"
            "- Lessons = future improvement"
        )
