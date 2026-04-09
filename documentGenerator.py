import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- PAGE CONFIG ---------------- #
st.set_page_config(page_title="Closure Document Generator", page_icon="📄", layout="wide")

# ---------------- STYLING ---------------- #
st.markdown("""
<style>
.block-container {padding-top: 2rem;}
h1, h2, h3 {color: #1f4e79;}
</style>
""", unsafe_allow_html=True)

# ---------------- SIDEBAR ---------------- #
with st.sidebar:
    st.header("📘 Smart Tips")
    st.info("""
✔ Write outcomes, not activities  
✔ Use numbers (%, count, time saved)  
✔ Keep it concise (decision makers skim)  
✔ Add proof via files/screenshots  
✔ Lessons = future improvement  
""")

# ---------------- HELPERS ---------------- #
def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    para = doc.add_paragraph(text if text else "N/A")
    para.paragraph_format.space_after = Pt(10)

def add_table(doc, data):
    table = doc.add_table(rows=1, cols=len(data[0]))
    for i, heading in enumerate(data[0]):
        table.rows[0].cells[i].text = heading
    for row in data[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_end)

def sanitize_filename(name):
    return "".join(c for c in name if c.isalnum() or c in (" ", "_", "-")).rstrip()

# ---------------- DELIVERABLES ---------------- #
def add_deliverables(doc, text, files):
    doc.add_heading("4. Deliverables", 1)
    add_paragraph(doc, text)

    if not files:
        return

    doc.add_paragraph("Attached Deliverables:")

    table_data = [["File Name", "Type"]]

    for file in files:
        file.seek(0)
        try:
            if file.type.startswith("image"):
                doc.add_paragraph(f"Image: {file.name}")
                doc.add_picture(BytesIO(file.read()), width=Inches(5))
            else:
                table_data.append([file.name, file.type])
        except:
            table_data.append([file.name, "Unsupported"])

    if len(table_data) > 1:
        add_table(doc, table_data)

# ---------------- CUSTOM SECTIONS ---------------- #
def add_custom_sections(doc, sections):
    if not sections:
        return

    doc.add_heading("8. Additional Sections", 1)

    for i, sec in enumerate(sections, start=1):
        if sec["title"]:
            doc.add_heading(f"{i}. {sec['title']}", 2)
            add_paragraph(doc, sec["content"])

# ---------------- DOCUMENT ---------------- #
def generate_document(data):
    doc = Document()

    doc.add_heading(data["project_name"], 0)
    add_paragraph(doc, "Project Closure Report")
    add_paragraph(doc, f"Generated on: {datetime.now().strftime('%d %b %Y')}")

    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    add_table_of_contents(doc)
    doc.add_page_break()

    sections = [
        ("1. Project Overview", data["overview"]),
        ("2. Objectives", data["objectives"]),
        ("3. Timeline", data["timeline"]),
    ]

    for title, content in sections:
        add_heading(doc, title, 1)
        add_paragraph(doc, content)

    add_deliverables(doc, data["deliverables"], data["files"])

    add_heading(doc, "5. Risks & Issues", 1)
    add_paragraph(doc, data["risks"])

    add_heading(doc, "6. Lessons Learned", 1)
    add_paragraph(doc, data["lessons"])

    add_heading(doc, "7. Stakeholders", 1)
    table = [["Name", "Role"]]
    table.extend(data["stakeholders"] or [["N/A", "N/A"]])
    add_table(doc, table)

    add_custom_sections(doc, data["custom_sections"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------- UI ---------------- #
st.title("📄 Project Closure Generator")

# ---------------- PROJECT INFO ---------------- #
with st.expander("📌 Project Information", expanded=True):
    st.caption("👉 Basic identification details")

    project_name = st.text_input(
        "Project Name *",
        help="Use official project name",
        placeholder="Migration of On-Prem Email to Cloud"
    )

    timeline = st.text_input(
        "Timeline",
        help="Project duration",
        placeholder="Jan 2025 – Mar 2025"
    )

# ---------------- OVERVIEW ---------------- #
with st.expander("📖 Project Overview", expanded=True):
    st.caption("👉 What was done, why, and outcome")

    st.info("Tip: Cover scope + purpose + final outcome")

    overview = st.text_area(
        "Overview *",
        help="Write a concise 3–5 line summary",
        placeholder="""Migrated 2500+ users to cloud, improving availability and reducing infra dependency."""
    )

# ---------------- OBJECTIVES ---------------- #
with st.expander("🎯 Objectives"):
    st.caption("👉 What success looked like")

    st.info("Tip: Use measurable goals")

    objectives = st.text_area(
        "Objectives *",
        placeholder="""- Reduce downtime < 1 hour
- Improve remote access
- Decommission legacy infra"""
    )

# ---------------- DELIVERABLES ---------------- #
with st.expander("📦 Deliverables"):
    st.caption("👉 What was delivered")

    st.info("Tip: Think outputs, not activities")

    deliverables = st.text_area(
        "Deliverables",
        placeholder="""- Mailboxes migrated
- Documentation created
- Dashboard deployed"""
    )

    uploaded_files = st.file_uploader(
        "Upload Supporting Files",
        accept_multiple_files=True,
        help="Upload screenshots, reports, dashboards"
    )

# ---------------- RISKS ---------------- #
with st.expander("⚠️ Risks & Issues"):
    st.caption("👉 What went wrong or could have")

    st.info("Tip: Mention impact + mitigation")

    risks = st.text_area(
        "Risks",
        placeholder="Data corruption risk, delays, etc."
    )

# ---------------- LESSONS ---------------- #
with st.expander("📘 Lessons Learned"):
    st.caption("👉 What should improve next time")

    st.info("Tip: Focus on actionable improvements")

    lessons = st.text_area(
        "Lessons",
        placeholder="Validate data earlier, improve planning buffer"
    )

# ---------------- CUSTOM SECTIONS ---------------- #
st.subheader("➕ Custom Sections")

if "custom_sections" not in st.session_state:
    st.session_state.custom_sections = []

if st.button("Add Section"):
    st.session_state.custom_sections.append({"title": "", "content": ""})

for i, sec in enumerate(st.session_state.custom_sections):
    st.session_state.custom_sections[i]["title"] = st.text_input(
        f"Title {i+1}",
        placeholder="E.g., Migration Strategy"
    )
    st.session_state.custom_sections[i]["content"] = st.text_area(
        f"Content {i+1}",
        placeholder="Describe this section"
    )

# ---------------- STAKEHOLDERS ---------------- #
with st.expander("👥 Stakeholders"):
    st.caption("👉 Who was involved")

    stakeholders = []
    count = st.number_input("Number of Stakeholders", 1, 10, 2)

    for i in range(count):
        name = st.text_input(f"Name {i+1}", key=f"name_{i}")
        role = st.text_input(f"Role {i+1}", key=f"role_{i}")
        if name and role:
            stakeholders.append([name, role])

# ---------------- VALIDATION ---------------- #
def validate():
    if not project_name.strip():
        st.error("Project Name required")
        return False
    if not overview.strip():
        st.error("Overview required")
        return False
    if not objectives.strip():
        st.error("Objectives required")
        return False
    return True


#-----------------Preview-----------------#

st.warning("⚠️ Review all sections carefully before generating the final document.")

st.divider()
st.subheader("📄 Document Preview")

if project_name:
    st.markdown(f"# {project_name}")

st.markdown(f"**Timeline:** {timeline or 'N/A'}")

with st.expander("1. Project Overview", expanded=True):
    st.write(overview or "N/A")

with st.expander("2. Objectives"):
    st.write(objectives or "N/A")

with st.expander("3. Timeline"):
    st.write(timeline or "N/A")

with st.expander("4. Deliverables"):
    st.write(deliverables or "N/A")

    if uploaded_files:
        st.markdown("**Attached Files:**")
        for f in uploaded_files:
            if f.type.startswith("image"):
                st.image(f, caption=f.name)
            else:
                st.write(f"📄 {f.name}")

with st.expander("5. Risks & Issues"):
    st.write(risks or "N/A")

with st.expander("6. Lessons Learned"):
    st.write(lessons or "N/A")

with st.expander("7. Stakeholders"):
    if stakeholders:
        for s in stakeholders:
            st.write(f"- {s[0]} ({s[1]})")
    else:
        st.write("N/A")

# Custom Sections Preview
if st.session_state.custom_sections:
    st.subheader("📌 Additional Sections")
    for i, sec in enumerate(st.session_state.custom_sections, start=1):
        if sec["title"]:
            with st.expander(f"{i}. {sec['title']}"):
                st.write(sec["content"] or "N/A")

# ---------------- GENERATE ---------------- #
st.divider()

if st.button("🚀 Generate Document"):
    if not validate():
        st.stop()

    data = {
        "project_name": project_name,
        "overview": overview,
        "objectives": objectives,
        "timeline": timeline,
        "deliverables": deliverables,
        "files": uploaded_files,
        "risks": risks,
        "lessons": lessons,
        "stakeholders": stakeholders,
        "custom_sections": st.session_state.custom_sections
    }

    doc = generate_document(data)

    filename = f"{sanitize_filename(project_name)}_Closure_Report.docx"

    st.success("✅ Document generated successfully!")

    st.download_button("📥 Download Document", doc, file_name=filename)